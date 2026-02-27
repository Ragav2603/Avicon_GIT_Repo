"""Document text extraction — supports PDF, DOCX, XLSX, CSV, TXT, MD."""
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger("avicon.doc_extractor")

MAX_CHARS = 15000  # Max characters to extract per document


def extract_text(file_path: str) -> Optional[str]:
    """Extract readable text from a document file.

    Returns extracted text or None if extraction fails.
    """
    path = Path(file_path)
    if not path.exists():
        logger.warning(f"File not found: {file_path}")
        return None

    ext = path.suffix.lower()
    try:
        if ext == ".pdf":
            return _extract_pdf(path)
        elif ext in (".docx", ".doc"):
            return _extract_docx(path)
        elif ext in (".xlsx", ".xls"):
            return _extract_xlsx(path)
        elif ext == ".csv":
            return _extract_csv(path)
        elif ext in (".txt", ".md"):
            return _extract_text(path)
        elif ext == ".pptx":
            return _extract_text(path)  # fallback
        else:
            return _extract_text(path)
    except Exception as e:
        logger.error(f"Extraction failed for {file_path}: {e}")
        return None


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    texts = []
    total = 0
    for page in reader.pages:
        text = page.extract_text() or ""
        texts.append(text)
        total += len(text)
        if total > MAX_CHARS:
            break
    return "\n\n".join(texts)[:MAX_CHARS]


def _extract_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    texts = []
    total = 0
    for para in doc.paragraphs:
        texts.append(para.text)
        total += len(para.text)
        if total > MAX_CHARS:
            break
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            texts.append(row_text)
            total += len(row_text)
            if total > MAX_CHARS:
                break
    return "\n".join(texts)[:MAX_CHARS]


def _extract_xlsx(path: Path) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(str(path), read_only=True, data_only=True)
    texts = []
    total = 0
    for sheet in wb.sheetnames[:3]:  # First 3 sheets
        ws = wb[sheet]
        texts.append(f"[Sheet: {sheet}]")
        for row in ws.iter_rows(max_row=200, values_only=True):
            row_text = " | ".join(str(v) for v in row if v is not None)
            if row_text.strip():
                texts.append(row_text)
                total += len(row_text)
                if total > MAX_CHARS:
                    break
        if total > MAX_CHARS:
            break
    wb.close()
    return "\n".join(texts)[:MAX_CHARS]


def _extract_csv(path: Path) -> str:
    import csv
    texts = []
    total = 0
    with open(path, "r", errors="ignore") as f:
        reader = csv.reader(f)
        for row in reader:
            row_text = " | ".join(row)
            texts.append(row_text)
            total += len(row_text)
            if total > MAX_CHARS:
                break
    return "\n".join(texts)[:MAX_CHARS]


def _extract_text(path: Path) -> str:
    return path.read_text(errors="ignore")[:MAX_CHARS]
